import docx as dd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import ApatcherClass as ac


# разбиение списка файлов на две категории - бд и веб
def split_list_files(lmass):
    lm_db = []
    lm_web = []
    for item in lmass:
        if "flexy-" in item:
            lm_web.append(item)
        else:
            lm_db.append(item)

    lm_db.sort()
    lm_web.sort()

    return lm_db, lm_web


# Функция генерации update_log.docx
def generate_doc_upd_log(author_name="Default", dir_name="\\0000-00-00_00\\", date_d="01 января 2000 года",
                         list_patch=["1", "2"]):
    dir_name = "\\" + dir_name + "\\"
    # переменные текста
    p_1_pr_str = "Устанавливаются обновления, полученные " + date_d + " от сотрудника отдела внедрения ООО \"ПОТОК\" "
    p_1_pr_str_1 = author_name + " в виде архива(директории) " + dir_name + "."
    p_2_pr_str = "Дата и время установки выполнения обновлений"
    p_2_pr_str_1 = " __________________________."

    # работа с документом
    document = dd.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # C Т И Л И
    # для заголовка
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TitleStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    # основной
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('BodyStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    # доп.заголовок
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TitleStyle2', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    # шапка
    p_begin = document.add_paragraph()
    p_begin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_begin.add_run("Журнал установки обновлений", style='TitleStyle1').bold = True

    # 1 блок
    p_1 = document.add_paragraph()
    p_1.add_run(p_1_pr_str + p_1_pr_str_1)

    # 2 блок
    p_2 = document.add_paragraph()
    p_2.add_run(p_2_pr_str + p_2_pr_str_1)

    # 3 блок
    p_3 = document.add_paragraph()
    r_3 = p_3.add_run()
    r_3.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_3.add_text(" Тестовая среда")
    r_3.add_tab()
    r_3.add_tab()
    r_3.add_tab()
    r_3.add_tab()
    r_3.add_tab()
    r_3.add_tab()
    r_3.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_3.add_text(" Продуктивная среда")
    r_3.add_break()

    # 4 блок
    p_4 = document.add_paragraph()
    p_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_4.add_run("Выполненные мероприятия:", style='TitleStyle2').bold = True

    # 5 блок
    p_5 = document.add_paragraph()
    r_5 = p_5.add_run()
    r_5.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_5.add_text(" Резервная копия сделана")
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_5.add_text(" Weblogic остановлен")
    r_5.add_break()
    r_5.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_5.add_text(" Задания Oracle остановлены")
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_tab()
    r_5.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_5.add_text(" Сборка обновлена")
    r_5.add_break()
    r_5.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_5.add_text(" Библиотеки обновлены")
    r_5.add_break()

    # 6 блок - таблица
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№п.п"
    hdr_cells[1].text = "Наименование обновления (патча)"
    hdr_cells[2].text = "Результат установки (Ок/Err)"
    hdr_cells[3].text = "Сотрудник, проводивший установку"
    hdr_cells[4].text = "Комментарий"
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].width = Cm(10)
    hdr_cells[2].width = Cm(1)
    hdr_cells[3].width = Cm(1)
    hdr_cells[4].width = Cm(7)
    run = hdr_cells[0].paragraphs[0].runs[0]
    pv_counter = 1
    for item in list_patch:
        row_cells = table.add_row().cells
        row_cells[0].text = str(pv_counter)
        row_cells[1].text = item
        row_cells[2].text = ""
        row_cells[3].text = ""
        row_cells[4].text = ""
        pv_counter += 1

    # 7 блок
    p_7_0 = document.add_paragraph()
    r_7_0 = p_7_0.add_run()
    r_7_0.add_break()
    p_7 = document.add_paragraph()
    p_7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_7.add_run("Выполненные мероприятия:", style='TitleStyle2').bold = True

    # 8 блок
    p_8 = document.add_paragraph()
    r_8 = p_8.add_run()
    r_8.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_8.add_text(" Weblogic запущен")
    r_8.add_tab()
    r_8.add_tab()
    r_8.add_tab()
    r_8.add_tab()
    r_8.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_8.add_text(" Задания Oracle запущены")
    r_8.add_break()
    r_8.add_picture("tmp/box_img.jpg", width=Inches(0.2))
    r_8.add_text(" Журналы weblogic проверены на наличие ошибок")
    r_8.add_break()

    # 9 блок
    p_9 = document.add_paragraph()
    p_9.add_run("Установку выполнил", style='BodyStyle1').bold = True
    r_9 = p_9.add_run()
    r_9.add_tab()
    r_9.add_tab()
    r_9.add_tab()
    r_9.add_tab()
    r_9.add_text(" ___________/_________________/")

    margin = 1.5
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin + 0.5)
        section.bottom_margin = Cm(margin + 0.5)
        section.left_margin = Cm(2 * margin)
        section.right_margin = Cm(margin)

    document.save('tmp//update_log.docx')


# Функция генерации changelist.docx
def generate_doc_changelist(project_patches, base_pacthes=[""], sdk_patches=[""], doc_upd=[""]):
    # работа с документом
    document = dd.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # C Т И Л И
    # для заголовка
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TitleStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    # основной
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('BodyStyle1', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
    # доп.заголовок
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TitleStyle2', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    # шапка
    p_begin = document.add_paragraph()
    p_begin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_begin.add_run("Описание изменений", style='TitleStyle1').bold = True

    # 1 блок
    p_1 = document.add_paragraph()
    r_1 = p_1.add_run()
    r_1.add_break()
    p_1.add_run("Таблица №1. Описание обновлений (Базовые патчи)", style='TitleStyle2').bold = True

    # 2 блок - таблица Базовых патчей
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№ патча"
    hdr_cells[1].text = "Краткое описание обновления (патча)"
    hdr_cells[2].text = "Перечень изменяемых объектов в БД"
    hdr_cells[3].text = "Перечень изменяемых объектов в веб-интерфейсе"
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].width = Cm(10)
    hdr_cells[2].width = Cm(10)
    hdr_cells[3].width = Cm(7)
    run = hdr_cells[0].paragraphs[0].runs[0]
    pv_counter = 1
    for item in base_pacthes:
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = item
        row_cells[2].text = ""
        row_cells[3].text = ""
        pv_counter += 1

    # 3 блок
    p_3 = document.add_paragraph()
    r_3 = p_3.add_run()
    r_3.add_break()
    p_3.add_run("Таблица №2. Описание обновлений (Патчи SDK)", style='TitleStyle2').bold = True

    # 4 блок - таблица SDK патчей
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№ патча"
    hdr_cells[1].text = "Краткое описание обновления (патча)"
    hdr_cells[2].text = "Перечень изменяемых объектов в БД"
    hdr_cells[3].text = "Перечень изменяемых объектов в веб-интерфейсе"
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].width = Cm(10)
    hdr_cells[2].width = Cm(10)
    hdr_cells[3].width = Cm(7)
    run = hdr_cells[0].paragraphs[0].runs[0]
    pv_counter = 1
    for item in sdk_patches:
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = item
        row_cells[2].text = ""
        row_cells[3].text = ""
        pv_counter += 1

    # 5 блок
    p_4 = document.add_paragraph()
    r_4 = p_4.add_run()
    r_4.add_break()
    p_4.add_run("Таблица №3. Описание обновлений (Проектные патчи)", style='TitleStyle2').bold = True

    # 6 блок - таблица проектных патчей
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№ патча"
    hdr_cells[1].text = "Краткое описание обновления (патча)"
    hdr_cells[2].text = "Перечень изменяемых объектов в БД"
    hdr_cells[3].text = "Перечень изменяемых объектов в веб-интерфейсе"
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].width = Cm(10)
    hdr_cells[2].width = Cm(10)
    hdr_cells[3].width = Cm(7)
    run = hdr_cells[0].paragraphs[0].runs[0]
    pv_counter = 1
    lm_db, lm_web = split_list_files(project_patches.list_files)
    # for item in project_patches:
    name_db = ",\n".join(lm_db)
    name_web = ",\n".join(lm_web)
    row_cells = table.add_row().cells
    row_cells[0].text = str(int(project_patches.name.split(".", 1)[0]))
    row_cells[1].text = project_patches.description
    row_cells[2].text = name_db
    row_cells[3].text = name_web
    pv_counter += 1

    # 7 блок
    p_7 = document.add_paragraph()
    r_7 = p_7.add_run()
    r_7.add_break()
    p_7.add_run("Таблица №4. Перечень обновляемой документации", style='TitleStyle2').bold = True

    # 8 блок - таблица проектных патчей
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Наименование документа"
    hdr_cells[2].text = "Новая версия документа"
    hdr_cells[0].width = Cm(3)
    hdr_cells[1].width = Cm(20)
    hdr_cells[2].width = Cm(10)
    run = hdr_cells[0].paragraphs[0].runs[0]
    pv_counter = 1
    for item in doc_upd:
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = item
        row_cells[2].text = ""
        pv_counter += 1

    margin = 1.5
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin + 0.5)
        section.bottom_margin = Cm(margin + 0.5)
        section.left_margin = Cm(2 * margin)
        section.right_margin = Cm(margin)

    document.save('tmp//changelist.docx')


def main():
    # generate_doc_upd_log("Богомолова А.В", "\\2016-01-01_05\\", "16 мая 2016 года", list_patch=["1.txt", "2.txt"])
    generate_doc_changelist(base_pacthes=["m1", "m2"], sdk_patches=[], project_patches=ac.PatchPrint(name="55",
                                                                                                     list_files=[
                                                                                                         "flexy-515100.sql",
                                                                                                         "pack_1.pck",
                                                                                                         "flexy-51000.sql",
                                                                                                         "pack_2.pck",
                                                                                                         "pack_3.pck"],
                                                                                                     description=""))


if __name__ == "__main__":
    main()
